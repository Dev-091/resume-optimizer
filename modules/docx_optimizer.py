import os
import tempfile
from pdf2docx import parse
from docx import Document

def replace_paragraph_text(paragraph, new_text):
    """Replace text in a paragraph while attempting to preserve the first run's formatting."""
    if len(paragraph.runs) == 0:
        paragraph.text = new_text
        return
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ""

def optimize_docx_resume(pdf_bytes, ats_keywords, allowed_skills, client):
    """
    Converts a PDF resume to DOCX, optimizes its bullet points
    in-place via LLM, and returns the path to the modified DOCX.
    """
    # Write pdf bytes to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
        tmp_pdf.write(pdf_bytes)
        pdf_path = tmp_pdf.name
        
    docx_path = pdf_path.replace('.pdf', '.docx')
    
    # Convert PDF to DOCX
    try:
        parse(pdf_path, docx_path)
    except Exception as e:
        # Fallback if parsing completely fails
        raise RuntimeError(f"Failed to convert PDF to Word document structure: {e}")
    
    # Read the DOCX
    doc = Document(docx_path)
    
    def process_text_block(old_text):
        old_text_clean = old_text.strip()
        # Only process substantial blocks (e.g., bullet points or summaries)
        # Short phrases are likely names, headers, dates, or small details.
        if len(old_text_clean.split()) < 8:
            return old_text
            
        prompt = f"""
        You are an expert ATS optimization engine. Your goal is to rewrite the following block of text from a resume to maximize its ATS score while preserving the user's original layout and formatting.
        
        CRITICAL RULES:
        1. CONTENT: Rewrite the text to naturally include:
           ATS Keywords: {', '.join(ats_keywords[:15])}
           Verified Skills: {', '.join(allowed_skills[:15])}
        2. STRUCTURE: Return ONLY the rewritten text. NO markdown (no **, no bullet points unless they were original), NO quotes, NO conversational filler.
        3. CONTEXT: If this is a header, company name, or date range, return it EXACTLY as provided. 
        4. BULLETS: If the original text starts with a bullet character (•, -, *), ensure the rewritten text starts with the SAME character. 
        5. LENGTH: Keep the character count within +/- 15% of the original to maintain the document's layout.
        
        Original Text:
        {old_text}
        """
        
        try:
            response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {"role": "system", "content": "Return only the rewritten text, nothing else."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.4
            )
            new_text = response.choices[0].message.content.strip()
            # Clean up potential LLM quoting
            if new_text.startswith('"') and new_text.endswith('"'):
                new_text = new_text[1:-1]
            return new_text
        except Exception:
            # On API error, return original
            return old_text

    # Iterate paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            new_text = process_text_block(paragraph.text.strip())
            if new_text != paragraph.text.strip() and new_text:
                replace_paragraph_text(paragraph, new_text)

    # Tables are commonly used in PDF-to-DOCX conversions for layout
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text and paragraph.text.strip():
                        new_text = process_text_block(paragraph.text.strip())
                        if new_text != paragraph.text.strip() and new_text:
                            replace_paragraph_text(paragraph, new_text)

    # Save the optimized DOCX
    optimized_docx_path = docx_path.replace('.docx', '_optimized.docx')
    doc.save(optimized_docx_path)
    
    # Clean up temp base files
    if os.path.exists(pdf_path):
        os.unlink(pdf_path)
    if os.path.exists(docx_path):
        os.unlink(docx_path)
        
    return optimized_docx_path
